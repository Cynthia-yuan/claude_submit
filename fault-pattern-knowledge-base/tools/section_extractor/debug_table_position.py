#!/usr/bin/env python3
"""
表格位置调试脚本
帮助诊断表格提取问题
"""

import sys
from docx import Document

def debug_document_structure(file_path, section_id):
    """打印文档结构，帮助调试表格位置"""
    doc = Document(file_path)

    print(f"=== 文档结构分析 ===\n")
    print(f"文件: {file_path}")
    print(f"查找章节: {section_id}\n")

    # 1. 找到章节的段落位置
    import re
    section_pattern = re.compile(r'^(\d+(?:\.\d+)*)\s+(.+)')

    print("=== 段落列表 (doc.paragraphs) ===")
    for idx, para in enumerate(doc.paragraphs[:20]):  # 只显示前20个
        text = para.text.strip()
        match = section_pattern.match(text)
        if match:
            print(f"段落{idx}: [章节] {match.group(1)} - {match.group(2)}")
        elif text:
            print(f"段落{idx}: {text[:50]}")

    # 2. 分析 element.body 结构
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P

    print("\n=== 元素序列 (doc.element.body) ===")
    para_count = 0  # 对应 doc.paragraphs 的索引
    table_count = 0

    for elem_idx, element in enumerate(doc.element.body[:30]):  # 只显示前30个
        if isinstance(element, CT_P):
            # 获取段落文本
            try:
                para = doc.paragraphs[para_count]
                text = para.text.strip()
                match = section_pattern.match(text)
                if match:
                    print(f"元素{elem_idx}: 段落{para_count} -> [章节] {match.group(1)}")
                    if match.group(1) == section_id:
                        print(f"       >>> 找到目标章节开始位置: 段落索引={para_count}")
                elif text:
                    print(f"元素{elem_idx}: 段落{para_count} -> {text[:40]}")
                else:
                    print(f"元素{elem_idx}: 段落{para_count} -> [空段落]")
            except:
                print(f"元素{elem_idx}: 段落{para_count} -> [无法读取]")
            para_count += 1

        elif isinstance(element, CT_Tbl):
            # 表格位于 para_count - 1 和 para_count 之间
            # 即：在段落 para_count - 1 之后
            table_position = para_count - 1
            print(f"元素{elem_idx}: 表格{table_count} -> 在段落{table_position}之后 (当前para_count={para_count})")
            table_count += 1

    print(f"\n=== 统计 ===")
    print(f"总段落数: {len(doc.paragraphs)}")
    print(f"总表格数: {len(doc.tables)}")
    print(f"处理的段落: {para_count}")
    print(f"处理的表格: {table_count}")

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("用法: python3 debug_table_position.py <document.docx> <章节编号>")
        print("示例: python3 debug_table_position.py test.docx 5.1.2.1")
        sys.exit(1)

    debug_document_structure(sys.argv[1], sys.argv[2])
