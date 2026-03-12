#!/usr/bin/env python3
"""
章节信息提取工具

从 Word 文档中提取指定章节的内容，识别表格并分类输出
"""

import re
import sys
import logging
from typing import Dict, List, Optional
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("错误: 需要安装 python-docx")
    print("请运行: pip install python-docx")
    sys.exit(1)

logger = logging.getLogger(__name__)


class SectionExtractor:
    """章节提取器"""

    def __init__(self):
        """初始化提取器"""
        # 章节编号的正则表达式
        self.section_pattern = re.compile(r'^(\d+(?:\.\d+)*)\s+(.+)')

    def extract_from_file(
        self,
        file_path: str,
        section_id: str
    ) -> Optional[Dict]:
        """
        从 docx 文件中提取指定章节

        Args:
            file_path: Word 文档路径
            section_id: 章节编号，如 "5.1.2.1"

        Returns:
            章节信息字典
        """
        try:
            doc = Document(file_path)
        except Exception as e:
            logger.error(f"读取文件失败: {e}")
            return None

        # 查找章节开始位置
        start_idx = None
        section_title = None

        for idx, para in enumerate(doc.paragraphs):
            match = self.section_pattern.match(para.text.strip())
            if match:
                current_section_id = match.group(1)
                if current_section_id == section_id:
                    start_idx = idx
                    section_title = match.group(2).strip()
                    logger.info(f"找到章节 {section_id} 在第 {idx + 1} 段")
                    break

        if start_idx is None:
            logger.warning(f"未找到章节: {section_id}")
            return None

        # 查找章节结束位置（下一个同级或更高级章节）
        end_idx = self._find_section_end(doc.paragraphs, start_idx, section_id)

        # 提取表格（在章节范围内的表格）
        tables = self._extract_tables_from_doc(doc, start_idx, end_idx)

        # 分类数据
        categorized = self._categorize_items(tables)

        return {
            'section_id': section_id,
            'title': section_title,
            'tables': tables,
            'categorized': categorized
        }

    def _find_section_end(
        self,
        paragraphs,
        start_idx: int,
        section_id: str
    ) -> int:
        """查找章节结束位置"""
        current_level = len(section_id.split('.'))

        for idx in range(start_idx + 1, len(paragraphs)):
            para = paragraphs[idx]
            match = self.section_pattern.match(para.text.strip())

            if match:
                next_section_id = match.group(1)
                next_level = len(next_section_id.split('.'))

                # 遇到同级或更高级章节 -> 停止
                if next_level <= current_level:
                    logger.debug(f"章节 {section_id} 在第 {idx} 段结束 (遇到同级/上级章节 {next_section_id})")
                    return idx

                # 遇到直接子章节 -> 也停止（不包含子章节的内容）
                if next_level == current_level + 1:
                    logger.debug(f"章节 {section_id} 在第 {idx} 段结束 (遇到子章节 {next_section_id})")
                    return idx

        return len(paragraphs)

    def _extract_tables_from_doc(
        self,
        doc,
        start_idx: int,
        end_idx: int
    ) -> List[Dict]:
        """提取指定段落范围内的表格"""
        tables_data = []

        # 获取文档中所有元素（段落和表格）的顺序
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P

        elements = []
        for element in doc.element.body:
            if isinstance(element, CT_P):
                elements.append(('para', element))
            elif isinstance(element, CT_Tbl):
                elements.append(('table', element))

        # 遍历元素，找到在范围内的表格
        para_count = 0  # 已遇到的段落数量（对应 doc.paragraphs 的索引）
        table_index = 0  # 当前处理的表格索引（与 enumerate(doc.tables) 一致）

        logger.info(f"开始提取，范围: 段落索引 {start_idx} 到 {end_idx-1}")

        for elem_type, element in elements:
            if elem_type == 'para':
                # 段落元素：当前段落对应 doc.paragraphs[para_count]
                # 处理完后再递增
                logger.debug(f"处理段落 {para_count}")
                para_count += 1

            elif elem_type == 'table':
                # 表格元素：表格在 doc.element.body 中的位置
                # 它位于段落 para_count - 1 和段落 para_count 之间
                # 也就是说：表格在段落 para_count - 1 "之后"
                # 它的段落位置应该是 para_count - 1
                table_position = para_count - 1
                logger.debug(f"表格{table_index}在段落{table_position}之后，范围检查: {start_idx} <= {table_position} < {end_idx}")

                if start_idx <= table_position < end_idx:
                    # 提取表格
                    try:
                        table = doc.tables[table_index]

                        # 解析表头
                        headers = []
                        for cell in table.rows[0].cells:
                            headers.append(cell.text.strip())

                        # 解析数据行
                        rows = []
                        for row in table.rows[1:]:
                            row_data = []
                            for cell in row.cells:
                                row_data.append(cell.text.strip())
                            rows.append(row_data)

                        tables_data.append({
                            'index': len(tables_data),
                            'headers': headers,
                            'rows': rows
                        })

                        logger.debug(f"提取表格 {len(tables_data)}（在段落{table_position}之后）")

                    except Exception as e:
                        logger.warning(f"解析表格失败: {e}")

                table_index += 1

        logger.info(f"从章节范围内提取了 {len(tables_data)} 个表格")
        return tables_data

    def _categorize_items(self, tables: List[Dict]) -> Dict[str, List]:
        """对表格内容进行分类"""
        categorized = {
            '删除': [],
            '新增': [],
            '变更': []
        }

        for table in tables:
            headers = table['headers']
            rows = table['rows']

            # 查找类型列（可能的列名）
            type_col_idx = self._find_type_column(headers)

            if type_col_idx is not None:
                for row in rows:
                    if type_col_idx < len(row):
                        item_type = row[type_col_idx]

                        # 根据类型归类
                        if any(keyword in item_type for keyword in ['删除', '移除', '废弃']):
                            categorized['删除'].append(row)
                        elif any(keyword in item_type for keyword in ['新增', '添加', '引入']):
                            categorized['新增'].append(row)
                        elif any(keyword in item_type for keyword in ['修改', '变更', '更新']):
                            categorized['变更'].append(row)

        return categorized

    def _find_type_column(self, headers: List[str]) -> Optional[int]:
        """查找类型/变更类型列的索引"""
        keywords = ['类型', '变更类型', '分类', 'category', 'type']

        for idx, header in enumerate(headers):
            if any(keyword in header for keyword in keywords):
                return idx

        return None

    def print_categorized(self, categorized: Dict[str, List]) -> None:
        """打印分类结果"""
        print("=" * 80)
        print("差异分类结果")
        print("=" * 80)

        for category, items in categorized.items():
            if items:
                print(f"\n【{category}】共 {len(items)} 项")
                print("-" * 80)

                for i, item in enumerate(items, 1):
                    # 将一行数据用 | 连接显示
                    line = " | ".join(item)
                    print(f"{i:2d}. {line}")

        print("\n" + "=" * 80)
        print(f"总计: 删除 {len(categorized['删除'])} 项, "
              f"新增 {len(categorized['新增'])} 项, "
              f"变更 {len(categorized['变更'])} 项")
        print("=" * 80)


def main():
    """命令行入口"""
    import argparse

    parser = argparse.ArgumentParser(
        description='从 Word 文档中提取指定章节并分类输出'
    )
    parser.add_argument(
        'file',
        help='Word 文档路径 (.docx)'
    )
    parser.add_argument(
        'section',
        help='章节编号，如 5.1.2.1'
    )
    parser.add_argument(
        '-o', '--output',
        help='输出到文本文件'
    )
    parser.add_argument(
        '-v', '--verbose',
        action='store_true',
        help='显示详细信息'
    )

    args = parser.parse_args()

    # 配置日志
    level = logging.DEBUG if args.verbose else logging.INFO
    logging.basicConfig(
        level=level,
        format='%(asctime)s - %(levelname)s - %(message)s'
    )

    # 提取章节
    extractor = SectionExtractor()
    result = extractor.extract_from_file(args.file, args.section)

    if result:
        # 检查是否提取到表格
        total_items = sum(len(items) for items in result['categorized'].values())

        if total_items == 0:
            print(f"\n⚠️  警告: 章节 {args.section} 中未找到任何差异项")
            print(f"\n诊断信息:")
            print(f"  章节标题: {result['title']}")
            print(f"  表格数量: {len(result['tables'])}")
            print(f"\n可能的原因:")
            print(f"  1. 表格中没有'类型'或'变更类型'列")
            print(f"  2. 类型列中的内容不匹配关键词（删除/新增/变更）")
            print(f"  3. 使用 -v 参数查看详细日志以诊断问题")
            if args.verbose:
                print(f"\n详细表格信息:")
                for i, table in enumerate(result['tables']):
                    print(f"  表格 {i+1} 表头: {table['headers']}")
                    if table['rows']:
                        print(f"    第一行数据: {table['rows'][0]}")
            return 0

        # 打印分类结果
        if args.output:
            # 保存到文件
            import io
            from contextlib import redirect_stdout

            with open(args.output, 'w', encoding='utf-8') as f:
                with io.StringIO() as buf, redirect_stdout(buf):
                    extractor.print_categorized(result['categorized'])
                    f.write(buf.getvalue())
            print(f"\n结果已保存到: {args.output}")
        else:
            # 输出到终端
            extractor.print_categorized(result['categorized'])

        return 0
    else:
        print(f"错误: 未找到章节 {args.section}")
        return 1


if __name__ == '__main__':
    sys.exit(main())
