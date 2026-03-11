#!/usr/bin/env python3
"""
创建测试用的 Word 文档
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_test_document():
    """创建测试文档"""
    doc = Document()

    # 标题
    title = doc.add_heading('系统版本差异文档', 0)

    # 章节 5.1
    doc.add_heading('5.1 接口变更', level=1)
    doc.add_paragraph('本节描述系统接口的主要变更。')

    # 章节 5.1.1
    doc.add_heading('5.1.1 概述', level=2)
    doc.add_paragraph('接口变更的总体说明。')

    # 表格1 - 在5.1.1章节中
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '变更类型'
    table.rows[0].cells[1].text = '接口名称'
    table.rows[0].cells[2].text = '变更前'
    table.rows[0].cells[3].text = '变更后'

    table.rows[1].cells[0].text = '接口删除'
    table.rows[1].cells[1].text = 'old_func_a'
    table.rows[1].cells[2].text = 'void old_func_a()'
    table.rows[1].cells[3].text = '-'

    table.rows[2].cells[0].text = '接口修改'
    table.rows[2].cells[1].text = 'update_func'
    table.rows[2].cells[2].text = 'int update()'
    table.rows[2].cells[3].text = 'int update(int flags)'

    table.rows[3].cells[0].text = '接口新增'
    table.rows[3].cells[1].text = 'new_feature'
    table.rows[3].cells[2].text = '-'
    table.rows[3].cells[3].text = 'void new_feature()'

    # 章节 5.1.2.1
    doc.add_heading('5.1.2.1 接口差异说明', level=3)
    doc.add_paragraph('本章节记录了两个版本之间的核心接口差异。')

    # 表格2 - 在5.1.2.1章节中（这是我们想要提取的）
    table = doc.add_table(rows=4, cols=6)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '变更类型'
    table.rows[0].cells[1].text = '接口名称'
    table.rows[0].cells[2].text = '变更前签名'
    table.rows[0].cells[3].text = '变更后签名'
    table.rows[0].cells[4].text = '影响'
    table.rows[0].cells[5].text = '备注'

    table.rows[1].cells[0].text = '接口删除'
    table.rows[1].cells[1].text = 'core_api'
    table.rows[1].cells[2].text = 'void core_api()'
    table.rows[1].cells[3].text = '-'
    table.rows[1].cells[4].text = '核心功能'
    table.rows[1].cells[5].text = '已废弃'

    table.rows[2].cells[0].text = '接口修改'
    table.rows[2].cells[1].text = 'data_handler'
    table.rows[2].cells[2].text = 'void handler(data_t*)'
    table.rows[2].cells[3].text = 'void handler(data_t*, int mode)'
    table.rows[2].cells[4].text = '一般'
    table.rows[2].cells[5].text = '增加模式参数'

    table.rows[3].cells[0].text = '接口删除'
    table.rows[3].cells[1].text = 'legacy_util'
    table.rows[3].cells[2].text = 'int util()'
    table.rows[3].cells[3].text = '-'
    table.rows[3].cells[4].text = '低'
    table.rows[3].cells[5].text = '移除辅助函数'

    doc.add_paragraph('这些变更会影响现有系统的兼容性。')

    # 章节 5.1.2.2
    doc.add_heading('5.1.2.2 兼容性说明', level=3)
    doc.add_paragraph('后续版本将保持API兼容性...')

    # 表格3 - 在5.1.2.2章节中
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '版本'
    table.rows[0].cells[1].text = '兼容性'
    table.rows[0].cells[2].text = '说明'
    table.rows[1].cells[0].text = 'v1.0'
    table.rows[1].cells[1].text = '100%'
    table.rows[1].cells[2].text = '完全兼容'

    # 保存文档
    doc.save('test_document.docx')
    print('✓ 测试文档已创建: test_document.docx')
    print('\n文档结构:')
    print('  5.1 接口变更')
    print('    5.1.1 概述 -> 包含1个表格（3行数据）')
    print('    5.1.2.1 接口差异说明 -> 包含1个表格（3行数据）')
    print('    5.1.2.2 兼容性说明 -> 包含1个表格（1行数据）')
    print('\n建议测试:')
    print('  python3 section_extractor.py test_document.docx 5.1.2.1')
    print('\n预期输出:')
    print('  应该只提取5.1.2.1章节的表格（2个删除项，1个变更项）')

if __name__ == '__main__':
    create_test_document()
